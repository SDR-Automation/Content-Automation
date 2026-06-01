import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_HEADING = RGBColor(0x1D, 0x4E, 0xD8)
C_BODY    = RGBColor(0x37, 0x41, 0x51)
C_MUTED   = RGBColor(0x6B, 0x72, 0x80)
C_GREEN   = RGBColor(0x10, 0xB9, 0x81)
C_RED     = RGBColor(0xEF, 0x44, 0x44)
C_ACCENT  = RGBColor(0x00, 0xC2, 0xFF)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK    = RGBColor(0x0D, 0x11, 0x17)

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _add_run_colored(para, text, color, bold=False, size_pt=10):
    run = para.add_run(text)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    return run

def _add_cover(doc, metadata):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, "0D1117")
    p1 = cell.paragraphs[0]
    p1.paragraph_format.left_indent = Inches(0.2)
    p1.paragraph_format.space_before = Pt(10)
    _add_run_colored(p1, metadata["content_type"].upper(), C_ACCENT, bold=True, size_pt=9)
    p2 = cell.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.2)
    _add_run_colored(p2, metadata["topic"], C_WHITE, bold=True, size_pt=16)
    p3 = cell.add_paragraph()
    p3.paragraph_format.left_indent = Inches(0.2)
    p3.paragraph_format.space_after = Pt(10)
    _add_run_colored(p3, f"Audience: {metadata['audience']}  |  {metadata['timestamp']}", C_MUTED, size_pt=9)
    doc.add_paragraph()

def _add_qa_table(doc, review):
    score    = review.get("score", 0)
    approved = review.get("approved", False)
    feedback = review.get("feedback", "")
    criteria = review.get("criteria_scores", {})
    h = doc.add_heading("Quality Assurance Report", level=2)
    if h.runs:
        h.runs[0].font.color.rgb = C_HEADING
    labels = [
        ("technical_accuracy", "Technical Accuracy"),
        ("relevance", "Relevance & Value"),
        ("seo", "SEO Compliance"),
        ("brand_tone", "Brand & Tone"),
        ("quality", "Quality & Clarity"),
    ]
    rows = [("Criteria", "Score/20")] + \
           [(lbl, str(criteria.get(k, "-"))) for k, lbl in labels] + \
           [("Total Score", f"{score}/100"),
            ("Status", "APPROVED" if approved else "REVISED BY REVIEWER")]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        cells = tbl.rows[i].cells
        if i == 0:
            _set_cell_bg(cells[0], "0D1117")
            _set_cell_bg(cells[1], "0D1117")
            _add_run_colored(cells[0].paragraphs[0], label, C_WHITE, bold=True, size_pt=9)
            _add_run_colored(cells[1].paragraphs[0], value, C_WHITE, bold=True, size_pt=9)
        elif i == len(rows) - 1:
            bg = "F0FDF4" if approved else "FEF2F2"
            _set_cell_bg(cells[0], bg)
            _set_cell_bg(cells[1], bg)
            r0 = cells[0].paragraphs[0].add_run(label)
            r0.font.bold = True
            r0.font.size = Pt(9)
            r1 = cells[1].paragraphs[0].add_run(value)
            r1.font.bold = True
            r1.font.size = Pt(9)
            r1.font.color.rgb = C_GREEN if approved else C_RED
        else:
            _set_cell_bg(cells[0], "F9FAFB")
            _set_cell_bg(cells[1], "F9FAFB")
            r0 = cells[0].paragraphs[0].add_run(label)
            r0.font.size = Pt(9)
            r0.font.color.rgb = C_BODY
            r1 = cells[1].paragraphs[0].add_run(value)
            r1.font.size = Pt(9)
            r1.font.color.rgb = C_BODY
    if feedback:
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(6)
        r = note.add_run(f"Reviewer Notes: {feedback[:400]}")
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = C_MUTED
    doc.add_paragraph()

def _add_content(doc, content):
    for line in content.split("\n"):
        if line.startswith("### "):
            h = doc.add_heading(line[4:], level=3)
            if h.runs:
                h.runs[0].font.size = Pt(11)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
            if h.runs:
                h.runs[0].font.color.rgb = C_HEADING
        elif line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            if h.runs:
                h.runs[0].font.color.rgb = C_HEADING
        elif line.startswith(("- ", "* ")):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line[2:])
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = C_BODY
        elif not line.strip():
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = C_BODY

def export_docx(content, metadata, review, path):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)
    _add_cover(doc, metadata)
    _add_qa_table(doc, review)
    _add_content(doc, content)
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run("LLM-1: Groq Llama 3.3 70B (Generator)  |  LLM-2: Claude Sonnet (Reviewer)")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = C_MUTED
    doc.save(path)